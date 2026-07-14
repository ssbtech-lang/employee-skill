# import pdfplumber
# import spacy
# import json
# import sys
# import os
# from docx import Document
# from spacy.matcher import PhraseMatcher

# nlp = spacy.load("en_core_web_sm")

# with open("skills.json", "r") as f:
#     skill_data = json.load(f)

# master_skills = []

# for role, skills in skill_data.items():
#     master_skills.extend(skills)

# master_skills = list(set(master_skills))


# with open("aliases.json", "r") as f:
#     aliases = json.load(f)

# matcher = PhraseMatcher(nlp.vocab, attr="LOWER")

# all_patterns = master_skills + list(aliases.keys())
# patterns = [nlp.make_doc(skill) for skill in all_patterns]

# matcher.add("SKILLS", patterns)

# # for pdf

# def extract_text_from_pdf(file_path):
#     text = ""

#     with pdfplumber.open(file_path) as pdf:
#         for page in pdf.pages:
#             page_text = page.extract_text()

#             if page_text:
#                 text += page_text + "\n"

#     return text


# # for docx

# def extract_text_from_docx(file_path):
#     doc = Document(file_path)

#     text = ""

#     for para in doc.paragraphs:
#         text += para.text + "\n"

#     return text

# # Extract basic info

# def extract_basic_info(text):
#     name = ""
#     email = ""

#     lines = text.split("\n")

#     for line in lines[:5]:
#         line = line.strip()

#         if line and "@" not in line and len(line.split()) <= 4:
#             name = line
#             break

#     for token in text.split():
#         if "@" in token and "." in token:
#             email = token
#             break

#     return name, email

# # Extract skills
# def extract_skills(text):
#     doc = nlp(text)

#     matches = matcher(doc)

#     found_skills = []

#     for match_id, start, end in matches:
#         skill = doc[start:end].text
#         found_skills.append(skill)

#     return list(set(found_skills))


# # Normalize skills
# def normalize_skills(skills):
#     normalized_skills = []

#     for skill in skills:
#         lower_skill = skill.lower().strip()

#         normalized = aliases.get(lower_skill, skill)

#         if normalized not in normalized_skills:
#             normalized_skills.append(normalized)

#     return normalized_skills



# def main():
#     if len(sys.argv) < 2:
#         print(json.dumps({"error": "No file path provided"}))
#         return

#     resume_path = sys.argv[1]

#     extension = os.path.splitext(resume_path)[1].lower()

#     if extension == ".pdf":
#         resume_text = extract_text_from_pdf(resume_path)

#     elif extension == ".docx":
#         resume_text = extract_text_from_docx(resume_path)

#     else:
#         print(json.dumps({"error": "Unsupported file type"}))
#         return

#     if not resume_text:
#         print(json.dumps({"error": "No text extracted"}))
#         return

#     name, email = extract_basic_info(resume_text)

#     extracted_skills = extract_skills(resume_text)

#     normalized_skills = normalize_skills(extracted_skills)

#     parsed_resume = {
#         "name": name,
#         "email": email,
#         "skills": normalized_skills,
#         "skillSource": {}
#     }

#     for skill in normalized_skills:
#         parsed_resume["skillSource"][skill] = "resume"

    
#     print(json.dumps(parsed_resume))


# if __name__ == "__main__":
#     main()








import pdfplumber
import spacy
import json
import sys
import os
import re
from docx import Document
from spacy.matcher import PhraseMatcher

nlp = spacy.load("en_core_web_sm")

with open("skills.json", "r") as f:
    skill_data = json.load(f)

master_skills = []

for role, skills in skill_data.items():
    master_skills.extend(skills)

master_skills = list(set(master_skills))

with open("aliases.json", "r") as f:
    aliases = json.load(f)

matcher = PhraseMatcher(nlp.vocab, attr="LOWER")

all_patterns = master_skills + list(aliases.keys())
patterns = [nlp.make_doc(skill) for skill in all_patterns]

matcher.add("SKILLS", patterns)

def extract_text_from_pdf(file_path):
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

def extract_basic_info(text):
    name = ""
    email = ""
    phone = ""
    experience = ""
    
    lines = text.split("\n")
    
    # Extract name (first few lines, not containing @ or numbers)
    for line in lines[:5]:
        line = line.strip()
        if line and "@" not in line and not re.search(r'\d', line) and len(line.split()) <= 4:
            name = line
            break
    
    # Extract email
    email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
    if email_match:
        email = email_match.group()
    
    # Extract phone
    phone_patterns = [
        r'\+?\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
        r'\d{3}[-.\s]?\d{3}[-.\s]?\d{4}',
        r'\(\d{3}\)\s?\d{3}[-.\s]?\d{4}'
    ]
    for pattern in phone_patterns:
        phone_match = re.search(pattern, text)
        if phone_match:
            phone = phone_match.group()
            break
    
    # Extract years of experience
    exp_patterns = [
        r'(\d+)\+?\s*(?:years?|yrs?)\s*(?:of)?\s*(?:experience|exp)',
        r'(?:experience|exp)\s*(?:of)?\s*(\d+)\+?\s*(?:years?|yrs?)'
    ]
    for pattern in exp_patterns:
        exp_match = re.search(pattern, text, re.IGNORECASE)
        if exp_match:
            experience = exp_match.group(0)
            break
    
    return name, email, phone, experience

def extract_skills(text):
    doc = nlp(text)
    matches = matcher(doc)
    found_skills = []
    
    for match_id, start, end in matches:
        skill = doc[start:end].text
        found_skills.append(skill)
    
    return list(set(found_skills))

def normalize_skills(skills):
    normalized_skills = []
    for skill in skills:
        lower_skill = skill.lower().strip()
        normalized = aliases.get(lower_skill, skill)
        if normalized not in normalized_skills:
            normalized_skills.append(normalized)
    return normalized_skills

def main():
    if len(sys.argv) < 2:
        print(json.dumps({"error": "No file path provided"}))
        return

    resume_path = sys.argv[1]
    extension = os.path.splitext(resume_path)[1].lower()

    if extension == ".pdf":
        resume_text = extract_text_from_pdf(resume_path)
    elif extension == ".docx":
        resume_text = extract_text_from_docx(resume_path)
    else:
        print(json.dumps({"error": "Unsupported file type"}))
        return

    if not resume_text:
        print(json.dumps({"error": "No text extracted"}))
        return

    name, email, phone, experience = extract_basic_info(resume_text)
    extracted_skills = extract_skills(resume_text)
    normalized_skills = normalize_skills(extracted_skills)

    parsed_resume = {
        "name": name,
        "email": email,
        "phone": phone,
        "experience": experience,
        "skills": normalized_skills,
        "skillSource": {}
    }

    for skill in normalized_skills:
        parsed_resume["skillSource"][skill] = "resume"
    
    print(json.dumps(parsed_resume))

if __name__ == "__main__":
    main()